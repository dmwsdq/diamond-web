"""Utility functions for DOCX template processing and placeholder replacement.

Two kinds of placeholders are supported:

  Simple placeholders  {{variable_name}}
      Replaced once with the matching value from the ``replacements`` dict.
      Keys in ``replacements`` must include the braces, e.g. ``'{{nama_ilap}}'``.

  Row placeholders     {{row.field_name}}
      Used in a table row to mark it as a *repeating template row*.
      The row is cloned once for every item in the ``row_data`` list and each
      ``{{row.field}}`` is replaced with ``item_dict['field']``.
      If ``row_data`` is empty the template row is silently removed.

Example template table::

    | No | Nama ILAP          | Jenis Data          | Periode          | Baris    |
    |----|--------------------|---------------------|------------------|----------|
    | {{row.no}} | {{row.nama_ilap}} | {{row.jenis_data}} | {{row.periode}} | {{row.baris}} |

Example row_data::

    row_data = [
        {'no': '1', 'nama_ilap': 'ILAP A', 'jenis_data': 'Bulanan', ...},
        {'no': '2', 'nama_ilap': 'ILAP B', 'jenis_data': 'Tahunan', ...},
    ]
"""

from copy import deepcopy
from io import BytesIO
from docx import Document
from docx.table import _Row as DocxRow


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _replace_in_paragraph(paragraph, replacements):
    """Replace all simple {{key}} placeholders in a paragraph's runs."""
    if not paragraph.runs:
        return
    full_text = ''.join(run.text for run in paragraph.runs)
    if not any(ph in full_text for ph in replacements):
        return
    new_text = full_text
    for placeholder, value in replacements.items():
        new_text = new_text.replace(placeholder, str(value if value is not None else '-'))
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)
    if new_text:
        paragraph.add_run(new_text)


def _row_has_row_placeholder(row):
    """Return True if any cell in *row* contains a ``{{row.xxx}}`` placeholder."""
    for cell in row.cells:
        for para in cell.paragraphs:
            if '{{row.' in ''.join(run.text for run in para.runs):
                return True
    return False


def _fill_row_placeholders(row, row_dict):
    """Replace ``{{row.field}}`` placeholders in all paragraphs of *row*."""
    for cell in row.cells:
        for para in cell.paragraphs:
            full_text = ''.join(run.text for run in para.runs)
            if '{{row.' not in full_text:
                continue
            new_text = full_text
            for key, value in row_dict.items():
                new_text = new_text.replace(
                    '{{row.' + key + '}}',
                    str(value if value is not None else '-'),
                )
            for run in para.runs:
                run._element.getparent().remove(run._element)
            if new_text:
                para.add_run(new_text)


def _expand_repeating_rows(table, row_data):
    """Expand every template row (one containing ``{{row.xxx}}``) in *table*.

    For each template row found:
    - If ``row_data`` is non-empty, clone the row once per item and fill
      ``{{row.field}}`` placeholders from the item dict.
    - The original template row is always removed afterwards.
    """
    # Collect indices of template rows (scan all rows once)
    template_indices = [
        i for i, row in enumerate(table.rows)
        if _row_has_row_placeholder(row)
    ]
    if not template_indices:
        return

    # Process from the last index to the first so earlier indices stay valid
    for idx in reversed(template_indices):
        template_tr = table.rows[idx]._element
        parent = template_tr.getparent()

        if row_data:
            # Insert cloned rows right after the template row, in order
            insert_after = template_tr
            for item_dict in row_data:
                new_tr = deepcopy(template_tr)
                insert_after.addnext(new_tr)
                insert_after = new_tr
                _fill_row_placeholders(DocxRow(new_tr, table), item_dict)

        # Remove the original template row
        parent.remove(template_tr)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def fill_template_with_data(template_file, replacements, row_data=None):
    """Fill a DOCX template with data by replacing placeholders.

    Args:
        template_file: File path string or file-like object of the template DOCX.
        replacements:  Dict of ``{'{{key}}': value}`` for simple placeholders.
        row_data:      Optional list of dicts for repeating table rows.
                       Each dict maps field names to values used in
                       ``{{row.field}}`` placeholders.

    Returns:
        BytesIO: The filled document ready for download.
    """
    doc = Document(template_file)

    # 1. Body paragraphs — simple replacements
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, replacements)

    # 2. Tables — first expand repeating rows, then simple replacements
    for table in doc.tables:
        _expand_repeating_rows(table, row_data)
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)

    # 3. Headers and footers — simple replacements only
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            _replace_in_paragraph(paragraph, replacements)
        for paragraph in section.footer.paragraphs:
            _replace_in_paragraph(paragraph, replacements)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output
